import streamlit as st
import json
import requests
import pdfplumber
import time
import os
import random
from datetime import datetime  # <-- NEW: Added for history timestamps
from io import BytesIO
from docx import Document
from dotenv import load_dotenv
import re
from bs4 import BeautifulSoup
from langchain_community.tools import DuckDuckGoSearchRun

# Load environment variables
load_dotenv()

# --- DUAL-KEY LOAD BALANCER ---
# Store original keys safely so we can rotate them dynamically
GROQ_KEYS = [os.environ.get("GROQ_API_KEY"), os.environ.get("GROQ_API_KEY_2")]
VALID_KEYS = [k for k in GROQ_KEYS if k]

def rotate_api_key():
    """Dynamically swaps the active API key to prevent Token-Per-Day limits."""
    if VALID_KEYS:
        os.environ["GROQ_API_KEY"] = random.choice(VALID_KEYS)

# Import our backend modules AFTER loading environment variables
from agents.profile_agent import extract_profile
from agents.module_a import run_module_a
from agents.module_b import run_module_b
from agents.module_c import run_module_c

# --- UI CONFIGURATION ---
st.set_page_config(page_title="Ascend AI | Student Navigator", page_icon="🎓", layout="wide")

# --- HELPER FUNCTIONS ---
@st.cache_data(ttl=86400) 
def get_live_exchange_rate():
    try:
        response = requests.get("https://open.er-api.com/v6/latest/USD", timeout=5)
        return response.json()["rates"]["PKR"]
    except:
        return 278.5 

def extract_text_from_pdf(uploaded_file):
    text = ""
    with pdfplumber.open(uploaded_file) as pdf:
        for page in pdf.pages:
            extracted = page.extract_text()
            if extracted: text += extracted + "\n"
    return text

def create_word_doc(text_content):
    doc = Document()
    doc.add_heading('Ascend AI - Document', 0)
    doc.add_paragraph(text_content)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- UPGRADED: The Jina AI Reader Bypass ---
def scrape_linkedin_public(url):
    try:
        jina_url = f"https://r.jina.ai/{url}"
        response = requests.get(jina_url, timeout=40)
        
        if response.status_code == 200:
            content = response.text
            if "authwall" in content.lower() or "sign in to view" in content.lower() or "join linkedin" in content.lower():
                return "LinkedIn Enterprise Firewall blocked the request (Authwall). Please use Option 2 (Upload CV) for accurate results."
            return f"LinkedIn Profile Data (via Jina Reader):\n{content[:3000]}"
        else:
            return f"Proxy blocked (Status {response.status_code}). Please use Option 2 (Upload CV)."
            
    except Exception as e:
        return f"Scraping Error: {str(e)}"

# --- INITIALIZE SESSION STATE ---
if "raw_notes" not in st.session_state:
    st.session_state.raw_notes = ""
if "generation_complete" not in st.session_state:
    st.session_state.generation_complete = False
if "profile_data" not in st.session_state:
    st.session_state.profile_data = None
if "module_a_data" not in st.session_state:
    st.session_state.module_a_data = None
if "module_b_data" not in st.session_state:
    st.session_state.module_b_data = None
if "module_c_data" not in st.session_state:
    st.session_state.module_c_data = None

# --- NEW: INITIALIZE HISTORY STATE ---
if "history" not in st.session_state:
    st.session_state.history = []

# --- UI RENDERING: SIDEBAR ---
st.sidebar.title("⚙️ System Status")
live_rate = get_live_exchange_rate()
st.sidebar.success(f"Live USD to PKR: **Rs. {live_rate:.2f}**")
st.sidebar.info("Architecture: Multi-Agent Load Balanced")

# --- NEW: SIDEBAR HISTORY UI ---
st.sidebar.divider()
st.sidebar.subheader("🗂️ Session History")

# Display a message if history is empty, otherwise show buttons
if not st.session_state.history:
    st.sidebar.caption("No previous searches yet. Generate a path to save it here!")
else:
    # Reverse the list so the newest generation is at the top
    for idx, record in enumerate(reversed(st.session_state.history)):
        # When clicked, it restores the saved data back into the main view
        if st.sidebar.button(f"📄 {record['title']}", key=f"hist_btn_{idx}", use_container_width=True):
            st.session_state.profile_data = record["profile_data"]
            st.session_state.module_a_data = record["module_a_data"]
            st.session_state.module_b_data = record["module_b_data"]
            st.session_state.module_c_data = record["module_c_data"]
            st.session_state.generation_complete = True
            st.rerun()

# --- MAIN UI ---
st.title("🎓 Ascend AI: Your Global & Local Career Navigator")

if st.button("✨ Load Demo Profile (Muhammad Kamal)"):
    st.session_state.raw_notes = (
        "Muhammad Kamal, BBA student at University of Sargodha (Malik Firoz Khan Noon Business School). "
        "Co-Founder & Lead Strategist at TriSyn Media. Founder of Skill Squad academic consultancy. "
        "Technical Lead at Lincoln Corner Sargodha and Social Media Manager for PAAP. "
        "Expertise in Python, SQL, and Power BI. Interests: Generative AI and Scholarship hunting for "
        "Master's abroad. Targets: Fulbright and Commonwealth scholarships."
    )
    st.session_state.generation_complete = False
    st.rerun()

# --- INPUT SECTION ---
col1, col2, col3 = st.columns([1.5, 1, 1])

with col1:
    user_text = st.text_area("Background & Goals:", key="raw_notes", height=200)

with col2:
    st.markdown("**Option 2: Upload CV**")
    uploaded_file = st.file_uploader("Standard CV", type=["pdf"])
    if uploaded_file:
        user_text = extract_text_from_pdf(uploaded_file)
        st.success("CV Parsed Successfully!")

with col3:
    st.markdown("**Option 3: LinkedIn URL**")
    li_url = st.text_input("Paste Public URL:", placeholder="linkedin.com/in/username")
    if st.button("🔍 Scrape Profile"):
        if li_url:
            with st.spinner("Scraping public data..."):
                scraped_data = scrape_linkedin_public(li_url)
                st.session_state.raw_notes = scraped_data
                st.session_state.generation_complete = False
                st.rerun()

# --- EXECUTION ENGINE ---
if st.button("🚀 Generate My Path", type="primary", use_container_width=True):
    if not user_text.strip():
        st.warning("Please provide background details first.")
    else:
        rotate_api_key()

        with st.status("Agent 1: Structuring Profile...", expanded=True) as status:
            st.session_state.profile_data = extract_profile(user_text)
            status.update(label="Profile Ready", state="complete", expanded=False)

        with st.status("Module A: Drafting SOP & Matching...", expanded=True) as status_a:
            st.session_state.module_a_data = run_module_a(st.session_state.profile_data, live_rate)
            status_a.update(label="SOP & Scholarships Ready!", state="complete", expanded=False)

        time.sleep(12)
        
        # --- NEW: Swap to Key 2 to bypass Groq rate limit ---
        rotate_api_key()

        with st.status("Module B: Mapping Careers...", expanded=True) as status_b:
            st.session_state.module_b_data = run_module_b(st.session_state.profile_data)
            status_b.update(label="Career Roadmap Ready!", state="complete", expanded=False)

        time.sleep(12)

        # --- NEW: Swap back to Key 1 (or keep Key 2) for safety ---
        rotate_api_key()

        with st.status("Module C: Local Tech Ecosystem...", expanded=True) as status_c:
            st.session_state.module_c_data = run_module_c(st.session_state.profile_data)
            status_c.update(label="Local Ecosystem Ready!", state="complete", expanded=False)
            status_c.update(label="Local Ecosystem Ready!", state="complete", expanded=False)

        # --- NEW: SAVE SNAPSHOT TO HISTORY BEFORE RERUN ---
        prof_name = st.session_state.profile_data.get('name', 'Student').split()[0] # Get first name
        timestamp = datetime.now().strftime("%H:%M") # Get current time
        
        st.session_state.history.append({
            "title": f"{prof_name} ({timestamp})",
            "profile_data": st.session_state.profile_data,
            "module_a_data": st.session_state.module_a_data,
            "module_b_data": st.session_state.module_b_data,
            "module_c_data": st.session_state.module_c_data
        })

        st.session_state.generation_complete = True
        st.rerun()

# --- UI RENDERING ---
if st.session_state.generation_complete:
    
    profile = st.session_state.profile_data
    module_a_results = st.session_state.module_a_data
    module_b_results = st.session_state.module_b_data
    module_c_results = st.session_state.module_c_data

    tab1, tab2, tab3, tab4 = st.tabs(["📝 Profile", "🌍 SOP & Scholarships", "💼 Career Roadmap", "🇵🇰 Local Ecosystem"])
    
    with tab1: 
        st.subheader(f"🎓 {profile.get('name', 'Student')} - Profile Overview")
        
        score = profile.get("profile_strength_score", 0)
        st.metric(label="Profile Strength Score", value=f"{score} / 100", delta="Competitiveness Matrix" if score > 75 else "Needs Improvement")
        st.progress(score / 100)
        
        st.divider()
        
        col_a, col_b = st.columns(2)
        with col_a:
            st.error("🚨 Gap Analysis")
            st.write(profile.get("gap_analysis", "No gaps identified."))
        with col_b:
            st.success("📈 6-Month Improvement Roadmap")
            for step in profile.get("improvement_roadmap", []):
                st.write(f"- {step}")
        
        st.divider()
        with st.expander("View Raw Extracted Data"):
            st.json(profile)

    with tab2:
        st.subheader("📝 Statement of Purpose")
        st.write(module_a_results["sop"])
            
        st.download_button("📄 Download SOP", create_word_doc(module_a_results["sop"]), "SOP.docx")
        
        st.divider()
        st.subheader("📊 Research Trends (Last 3 Years)")
        st.info("The hottest domains in your field based on recent ArXiv publications.")
        st.write(module_a_results.get("scholar_trends", ""))
        
        st.divider()
        st.subheader("🌍 Matched Scholarships (2026)")
        st.write(module_a_results["scholarships"])
        
        st.divider()
        st.subheader("🎯 Targeted Scholarship Gap Analysis")
        st.error("What you are specifically missing for the matched programs above.")
        st.write(module_a_results.get("targeted_gaps", ""))
        
        st.divider()
        st.subheader("🕵️ Live Deadline & Document Audit")
        st.info("Scraped directly from the official scholarship portals just now.")
        st.write(module_a_results["verification"])
        
        st.divider()
        st.subheader("👨‍🏫 Faculty Match & Cold Email Draft")
        st.info("Potential faculty advisors and a personalized cold email template referencing recent research.")
        st.write(module_a_results["faculty_outreach"])
        
        st.divider()
        st.subheader("🔬 AI-Drafted Research Proposal")
        st.info("A 500-word, literature-backed research proposal generated from live ArXiv papers.")
        st.write(module_a_results["proposal"])
        st.download_button("📄 Download Proposal", create_word_doc(module_a_results["proposal"]), "Research_Proposal.docx")

    with tab3:
        st.subheader("💼 Digital Career Mapping")
        st.write(module_b_results["careers"])
        
        st.divider()
        st.subheader("📺 12-Week YouTube Learning Roadmap")
        st.write(module_b_results["roadmap"])
        
        st.divider()
        st.subheader("💸 Freelance Market Entry")
        st.info("Ready-to-use Upwork bio, Fiverr gigs, and a 14-day portfolio project.")
        st.write(module_b_results["freelance"])
        
        st.divider()
        st.subheader("⚖️ Academia vs Industry (The Hard Truth)")
        st.info("An honest tradeoff analysis of pursuing a PhD vs entering the workforce.")
        st.write(module_b_results.get("phd_tradeoff", ""))

    with tab4:
        st.subheader("🇵🇰 Local Tech Ecosystem & Grants")
        st.info("Accelerators, funding opportunities, and communities tailored to your profile in Pakistan.")
        st.write(module_c_results["ecosystem"])
        
        st.divider()
        st.success("💡 Pro Tip: Connecting with your local National Incubation Center (NIC) or university Business Incubation Center is the fastest way to build your network before applying abroad.")